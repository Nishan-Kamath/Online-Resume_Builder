from flask import Flask, render_template, request, redirect, url_for, send_file
from docx import Document
import os
import io

app = Flask(__name__)

# Function to generate basic suggestions for resume improvement
def generate_suggestions(resume_data):
    suggestions = []

    # Example rule-based suggestions
    if 'python' not in resume_data['skills']:
        suggestions.append("Consider adding Python as a skill.")
    
    if len(resume_data['summary'].split()) < 50:
        suggestions.append("Your professional summary is too short. Aim for at least 50 words.")
    
    if len(resume_data['experience']) < 3:
        suggestions.append("You might want to add more details to your work experience section.")

    return suggestions

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/submit', methods=['POST'])
def submit_resume():
    # Get form data
    resume_data = {
        "name": request.form['name'],
        "email": request.form['email'],
        "phone": request.form['phone'],
        "summary": request.form['summary'],
        "skills": request.form.getlist('skills[]'),
        "experience": request.form['experience'],
    }

    # Generate suggestions
    suggestions = generate_suggestions(resume_data)

    return render_template('suggestions.html', resume_data=resume_data, suggestions=suggestions)

@app.route('/download_docx', methods=['POST'])
def download_docx():
    # Get resume data from the form
    resume_data = {
        "name": request.form['name'],
        "email": request.form['email'],
        "phone": request.form['phone'],
        "summary": request.form['summary'],
        "skills": request.form.getlist('skills[]'),
        "experience": request.form['experience'],
    }

    # Create a new Word document
    doc = Document()

    # Add content to the Word document
    doc.add_heading('Resume', 0)

    doc.add_heading('Name:', level=1)
    doc.add_paragraph(resume_data['name'])

    doc.add_heading('Email:', level=1)
    doc.add_paragraph(resume_data['email'])

    doc.add_heading('Phone:', level=1)
    doc.add_paragraph(resume_data['phone'])

    doc.add_heading('Professional Summary:', level=1)
    doc.add_paragraph(resume_data['summary'])

    doc.add_heading('Skills:', level=1)
    doc.add_paragraph(', '.join(resume_data['skills']))

    doc.add_heading('Experience:', level=1)
    doc.add_paragraph(resume_data['experience'])

    # Save the document to a BytesIO object
    docx_io = io.BytesIO()
    doc.save(docx_io)
    docx_io.seek(0)

    # Send the document as a response
    return send_file(docx_io, as_attachment=True, download_name='resume.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

if __name__ == '__main__':
    app.run(debug=True)
